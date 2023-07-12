import PyPDF2
import json
from docx import Document
from flask import request
import re


def pdf_to_json(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        num_pages = len(pdf_reader.pages)
        extracted_data = []

        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            extracted_data.append(page.extract_text())

        # Convert extracted data to JSON
        json_data = json.dumps(extracted_data)

        # Define the output file path
        json_output_path = 'path_to_output_json_file.json'

        # Write JSON data to the output file
        with open(json_output_path, 'w') as json_file:
            # json_file.write(json_data)
            return json_data
    except Exception as e:
        print("The Following Error occurred" + str(e))


def pdf():
    pdf_file = request.files['document']
    return pdf_to_json(pdf_file)


def doc_to_json(file_path):

    # document = Document('/Users/elisabethhassan/Documents/IntelliJProjects/ai dectector/English Research MLA2_.docx')
    document = Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]

    data = {
        'paragraphs': paragraphs,
    }
    json_result = json.dumps(data)
    res = " ".join(re.findall("[a-zA-Z]+", json_result[17:]))
    return res


def doc():
    doc_file = request.files['document']
    return doc_to_json(doc_file)